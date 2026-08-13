import os
import hashlib
import streamlit as st
from dotenv import load_dotenv

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_openai import OpenAIEmbeddings
try:
    from langchain_core.documents import Document
except Exception:
    # Fallback minimal Document for environments where langchain_core isn't available
    from dataclasses import dataclass

    @dataclass
    class Document:
        page_content: str
        metadata: dict

from PyPDF2 import PdfReader
import docx

# --- Load environment variables ---
load_dotenv()

# --- Embeddings ---
embeddings = OpenAIEmbeddings(
    model="text-embedding-3-small",
    api_key=os.getenv("OPENAI_API_KEY")
)

# --- Helpers ---
def extract_text(file_path: str) -> str:
    """Extract text from PDF, DOCX, or TXT files."""
    try:
        if file_path.endswith(".pdf"):
            reader = PdfReader(file_path)
            return "\n".join([page.extract_text() or "" for page in reader.pages])
        elif file_path.endswith(".docx"):
            doc = docx.Document(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        elif file_path.endswith(".txt"):
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
    except Exception as e:
        st.error(f"Error reading {file_path}: {e}")
    return ""

def chunk_document(content: str, source: str, file_path: str):
    """Split document into chunks and attach metadata for citations."""
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    docs = [
        Document(
            page_content=content,
            metadata={"source": source, "filepath": file_path}
        )
    ]
    return splitter.split_documents(docs)

def _hash_content(content: str) -> str:
    """Generate a SHA256 hash of the document content for duplicate detection."""
    return hashlib.sha256(content.encode("utf-8")).hexdigest()

# --- Load existing index ---
def load_vector_store(index_name="combined_ops_index"):
    """Load an existing FAISS index from disk."""
    save_path = os.path.join("data", index_name)
    if os.path.exists(save_path):
        return FAISS.load_local(save_path, embeddings, allow_dangerous_deserialization=True)
    return None

# --- Incremental update with duplicate detection ---
def update_vector_store(file_path, index_name=None):
    """
    Add a new file to FAISS index, skipping duplicates.
    Routes SOR files to 'SOR_index' and others to 'combined_ops_index'.
    """
    content = extract_text(file_path)
    if not content.strip():
        return None

    filename = os.path.basename(file_path).lower()
    index_name = "SOR_index" if "sor" in filename else (index_name or "combined_ops_index")

    chunks = chunk_document(content, os.path.basename(file_path), file_path)
    save_path = os.path.join("data", index_name)

    if os.path.exists(save_path):
        vector_store = FAISS.load_local(save_path, embeddings, allow_dangerous_deserialization=True)
        existing_sources = [doc.metadata.get("source") for doc in vector_store.docstore._dict.values()]
        if os.path.basename(file_path) in existing_sources:
            st.warning(f"⚠️ File already indexed, skipping: {os.path.basename(file_path)}")
            return vector_store
        vector_store.add_documents(chunks)
        st.success(f"✅ Indexed new file into {index_name}: {os.path.basename(file_path)}")
    else:
        vector_store = FAISS.from_documents(chunks, embeddings)
        st.success(f"✅ Created new index {index_name} with file: {os.path.basename(file_path)}")

    vector_store.save_local(save_path)
    return vector_store

# --- Build index from folder(s) ---
def build_vector_store(folders, index_name):
    """
    Rebuild FAISS index from all files in one or more folders.
    Args:
        folders: str or list of folder paths
        index_name: name of the FAISS index to save
    """
    if isinstance(folders, str):
        folders = [folders]

    all_chunks = []
    for folder in folders:
        if not os.path.exists(folder):
            continue
        for file in os.listdir(folder):
            file_path = os.path.join(folder, file)
            content = extract_text(file_path)
            if not content.strip():
                continue
            chunks = chunk_document(content, file, file_path)
            all_chunks.extend(chunks)

    if not all_chunks:
        st.warning(f"⚠️ No valid documents found in {folders}")
        return None

    save_path = os.path.join("data", index_name)
    vector_store = FAISS.from_documents(all_chunks, embeddings)
    vector_store.save_local(save_path)
    st.success(f"✅ Rebuilt index {index_name} with {len(all_chunks)} chunks")
    return vector_store
